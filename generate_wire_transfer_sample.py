import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new document
doc = docx.Document()

# Add a title
title = doc.add_heading('Wire Transfer Sample Information', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add a subtitle with the current date
doc.add_paragraph('For Testing and Validation Purposes Only', style='Subtitle')
doc.add_paragraph('May 14, 2025', style='Subtitle')

# Add source account section
doc.add_heading('Source Account', 1)
source = doc.add_paragraph('Treasury Account: ')
source.add_run('Emapa Pacific Fund - USD 1300001645').bold = True

# Add correspondent bank section
doc.add_heading('Correspondent Bank', 1)
doc.add_paragraph('Global Correspondent Bank (GCBA15333A)')

# Add transfer details section
doc.add_heading('Transfer Details', 1)
details_table = doc.add_table(rows=2, cols=2)
details_table.style = 'Table Grid'

# Fill in transfer details
cells = details_table.rows[0].cells
cells[0].text = 'Amount'
cells[1].text = 'USD 1000.00'
cells = details_table.rows[1].cells
cells[0].text = 'Purpose'
cells[1].text = 'Test wire transfer for system validation'

# Add originator information section
doc.add_heading('Originator Information', 1)
orig_table = doc.add_table(rows=3, cols=2)
orig_table.style = 'Table Grid'

# Fill in originator information
cells = orig_table.rows[0].cells
cells[0].text = 'Name'
cells[1].text = 'NVC Global Treasury'
cells = orig_table.rows[1].cells
cells[0].text = 'Account'
cells[1].text = '1300001645'
cells = orig_table.rows[2].cells
cells[0].text = 'Address'
cells[1].text = '123 Financial Plaza, Suite 400, Singapore 123456'

# Add beneficiary information section
doc.add_heading('Beneficiary Information', 1)
ben_table = doc.add_table(rows=3, cols=2)
ben_table.style = 'Table Grid'

# Fill in beneficiary information
cells = ben_table.rows[0].cells
cells[0].text = 'Name'
cells[1].text = 'Test Corporation Ltd.'
cells = ben_table.rows[1].cells
cells[0].text = 'Account'
cells[1].text = '9876543210'
cells = ben_table.rows[2].cells
cells[0].text = 'Address'
cells[1].text = '456 Commerce Street, London, UK EC4R 1AB'

# Add beneficiary bank information section
doc.add_heading('Beneficiary Bank Information', 1)
ben_bank_table = doc.add_table(rows=4, cols=2)
ben_bank_table.style = 'Table Grid'

# Fill in beneficiary bank information
cells = ben_bank_table.rows[0].cells
cells[0].text = 'Bank Name'
cells[1].text = 'International Trust Bank'
cells = ben_bank_table.rows[1].cells
cells[0].text = 'Bank Address'
cells[1].text = '789 Banking Boulevard, London, UK EC3M 4BY'
cells = ben_bank_table.rows[2].cells
cells[0].text = 'SWIFT/BIC Code'
cells[1].text = 'INTRGB2L'
cells = ben_bank_table.rows[3].cells
cells[0].text = 'Routing Number'
cells[1].text = '021000089 (for US banks)'

# Add intermediary bank information section
doc.add_heading('Intermediary Bank (Optional)', 1)
inter_bank_table = doc.add_table(rows=2, cols=2)
inter_bank_table.style = 'Table Grid'

# Fill in intermediary bank information
cells = inter_bank_table.rows[0].cells
cells[0].text = 'Bank Name'
cells[1].text = 'Global Clearing Bank'
cells = inter_bank_table.rows[1].cells
cells[0].text = 'SWIFT/BIC Code'
cells[1].text = 'GLCBUS33'

# Add additional information section
doc.add_heading('Additional Information', 1)
add_table = doc.add_table(rows=1, cols=2)
add_table.style = 'Table Grid'

# Fill in additional information
cells = add_table.rows[0].cells
cells[0].text = 'Message to Beneficiary'
cells[1].text = 'Payment for invoice #INV-2025-001'

# Add a footer note
doc.add_paragraph()
footer = doc.add_paragraph('This sample information is provided for testing the wire transfer system functionality only.')
footer.style = 'Intense Quote'

# Save the document
doc.save('wire_transfer_sample_information.docx')

print("Document 'wire_transfer_sample_information.docx' created successfully.")