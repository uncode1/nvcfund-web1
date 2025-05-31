"""
Generate a Word Document (.docx) Letter of Intent (LOI) for NVC Fund and Bridge.xyz partnership
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime

def generate_bridge_xyz_docx():
    """Generate a Word document for NVC Fund and Bridge.xyz Letter of Intent."""
    # Create output directory if it doesn't exist
    output_dir = os.path.join('static', 'docs')
    os.makedirs(output_dir, exist_ok=True)
    
    # Output file path
    docx_path = os.path.join(output_dir, 'NVCT_Bridge_XYZ_LOI.docx')
    
    # Create document
    doc = Document()
    
    # Set normal style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add title
    title = doc.add_heading('LETTER OF INTENT', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph('STRATEGIC PARTNERSHIP BETWEEN')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    
    subtitle2 = doc.add_paragraph('NVC FUND HOLDING TRUST AND BRIDGE.XYZ')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.runs[0].bold = True
    
    # Add date
    today = datetime.now().strftime("%B %d, %Y")
    date_paragraph = doc.add_paragraph(f'Date: {today}')
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add addressee
    doc.add_paragraph('Bridge.xyz')
    doc.add_paragraph('San Francisco, CA')
    doc.add_paragraph('United States')
    
    # Add salutation
    doc.add_paragraph('Re: Letter of Intent - $5 Billion Strategic Liquidity Partnership').bold = True
    doc.add_paragraph('Dear Bridge.xyz Leadership Team:')
    
    # Introduction paragraph
    intro_text = """
    This Letter of Intent ("LOI") sets forth the principal terms and conditions under which NVC Fund Holding Trust 
    ("NVC Fund") proposes to enter into a strategic liquidity partnership with Bridge.xyz ("Bridge") to establish 
    a $5 billion liquidity facility leveraging NVC Fund's NVCToken ("NVCT"). This non-binding LOI outlines the
    framework for our proposed partnership and serves as the basis for the preparation of definitive agreements.
    """
    doc.add_paragraph(intro_text)
    
    # 1. Partnership Structure
    doc.add_heading('1. PARTNERSHIP STRUCTURE', level=2)
    
    structure_text = """
    1.1 Strategic Liquidity Partnership: NVC Fund and Bridge.xyz propose to establish a multi-tiered liquidity 
    provision agreement with the following parameters:
    
    a) Initial Phase: $500 million liquidity provision against NVCToken collateral with a loan-to-value 
    ratio of 75%.
    
    b) Scaling Phases: Incremental increases to the full $5 billion facility based on predetermined 
    transaction volume and performance metrics over a 24-month period.
    
    c) Term: Initial 36-month arrangement with automatic renewal options upon mutual agreement.
    
    1.2 Technical Integration: The parties will collaborate to develop:
    
    a) Direct API connections between NVC Banking Platform and Bridge.xyz infrastructure
    
    b) Stablecoin settlement layer utilizing Bridge's existing stablecoin infrastructure
    
    c) Cross-chain bridge solutions to support NVCToken liquidity across multiple blockchain networks
    """
    doc.add_paragraph(structure_text)
    
    # 2. Commercial Terms
    doc.add_heading('2. COMMERCIAL TERMS', level=2)
    
    commercial_text = """
    2.1 Fee Structure: A tiered transaction fee model based on volume:
    
    a) Tier 1 (0-$500M): 0.40% per transaction
    
    b) Tier 2 ($500M-$2B): 0.30% per transaction
    
    c) Tier 3 ($2B-$5B): 0.25% per transaction
    
    2.2 Revenue Sharing: 70% to NVC Fund and 30% to Bridge.xyz for all new transaction revenue generated 
    through the partnership.
    
    2.3 Minimum Volume Commitments: 
    
    a) Year 1: $2 billion in transaction volume
    
    b) Year 2: $5 billion in transaction volume
    
    c) Year 3: $8 billion in transaction volume
    """
    doc.add_paragraph(commercial_text)
    
    # 3. Security and Compliance
    doc.add_heading('3. SECURITY AND COMPLIANCE', level=2)
    
    security_text = """
    3.1 Collateral Management: Establishment of a programmatic treasury management system with 
    transparent reporting and automated collateral adjustments.
    
    3.2 Smart Contract Governance: Formation of a joint oversight committee for contract modifications
    with representation from both parties.
    
    3.3 Regulatory Framework: Development of clear jurisdictional parameters and compliance requirements
    in accordance with relevant financial regulations and blockchain governance standards.
    """
    doc.add_paragraph(security_text)
    
    # 4. Growth Incentives
    doc.add_heading('4. GROWTH INCENTIVES', level=2)
    
    incentives_text = """
    4.1 Transaction Volume Bonuses: Automatic credit facility increases upon exceeding volume targets:
    
    a) 10% additional credit for exceeding quarterly targets by 15%
    
    b) 20% additional credit for exceeding quarterly targets by 25%
    
    4.2 Market Expansion: Preferential terms for entering new geographic markets, with specific focus on
    emerging markets in Africa, Southeast Asia, and Latin America.
    
    4.3 Technology Collaboration: Joint development of new liquidity products with intellectual property
    sharing arrangements to be detailed in the definitive agreements.
    """
    doc.add_paragraph(incentives_text)
    
    # 5. Mutual Benefits
    doc.add_heading('5. MUTUAL BENEFITS', level=2)
    
    benefits_text = """
    5.1 Benefits to NVC Fund:
    
    a) Immediate access to global stablecoin infrastructure
    
    b) Enhanced liquidity for NVCToken
    
    c) New distribution channels for NVC Fund services
    
    5.2 Benefits to Bridge.xyz:
    
    a) Access to NVC Fund's $10+ trillion asset base as backing collateral
    
    b) Enhanced credibility through partnership with an established financial institution
    
    c) Expanded transaction volume through NVC Fund's existing global networks
    """
    doc.add_paragraph(benefits_text)
    
    # 6. Non-Binding Nature
    doc.add_heading('6. NON-BINDING NATURE', level=2)
    
    nonbinding_text = """
    This LOI is non-binding and subject to the execution of definitive agreements between the parties.
    The parties agree to work in good faith to negotiate and execute such agreements within 60 days from
    the date of this LOI.
    """
    doc.add_paragraph(nonbinding_text)
    
    # 7. Confidentiality
    doc.add_heading('7. CONFIDENTIALITY', level=2)
    
    confidentiality_text = """
    The parties agree to maintain the confidentiality of this LOI and all discussions and information
    exchanged in connection with the proposed partnership, except as required by law or regulatory authorities.
    """
    doc.add_paragraph(confidentiality_text)
    
    # 8. Governing Law
    doc.add_heading('8. GOVERNING LAW', level=2)
    
    law_text = """
    This LOI shall be governed by and construed in accordance with the laws of the jurisdiction to be
    determined in the definitive agreements.
    """
    doc.add_paragraph(law_text)
    
    # Signature block
    doc.add_paragraph('If the terms outlined in this LOI are acceptable, please indicate your acceptance by signing below.')
    
    # Create signature table
    signature_table = doc.add_table(rows=6, cols=2)
    signature_table.style = 'Table Grid'
    
    # Header row
    header_cells = signature_table.rows[0].cells
    header_cells[0].text = 'FOR NVC FUND HOLDING TRUST:'
    header_cells[1].text = 'FOR BRIDGE.XYZ:'
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
    
    # Empty row for signatures
    signature_cells = signature_table.rows[1].cells
    signature_cells[0].text = ''
    signature_cells[1].text = ''
    
    # Signature line row
    sign_line_cells = signature_table.rows[2].cells
    sign_line_cells[0].text = '________________________________'
    sign_line_cells[1].text = '________________________________'
    
    # Name row
    name_cells = signature_table.rows[3].cells
    name_cells[0].text = 'Name: _________________________'
    name_cells[1].text = 'Name: _________________________'
    
    # Title row
    title_cells = signature_table.rows[4].cells
    title_cells[0].text = 'Title: __________________________'
    title_cells[1].text = 'Title: __________________________'
    
    # Date row
    date_cells = signature_table.rows[5].cells
    date_cells[0].text = 'Date: __________________________'
    date_cells[1].text = 'Date: __________________________'
    
    # Set table properties
    for row in signature_table.rows:
        row.height = Inches(0.5)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Save the document
    doc.save(docx_path)
    
    print(f"Generated LOI Word document at: {docx_path}")
    return docx_path

if __name__ == "__main__":
    generate_bridge_xyz_docx()