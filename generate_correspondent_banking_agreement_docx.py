"""
Script to generate a Correspondent Banking Agreement in Microsoft Word format for NVC Fund Bank.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

def generate_correspondent_banking_agreement_docx():
    """Generate a Word document for the NVC Fund Bank Correspondent Banking Agreement."""
    # Initialize document
    doc = Document()
    
    # Set document margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Define styles
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('Title Style', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(16)
    title_font.bold = True
    
    # Heading style
    heading_style = styles.add_style('Heading Style', WD_STYLE_TYPE.PARAGRAPH)
    heading_font = heading_style.font
    heading_font.name = 'Arial'
    heading_font.size = Pt(12)
    heading_font.bold = True
    
    # Normal text style
    normal_style = styles.add_style('Normal Style', WD_STYLE_TYPE.PARAGRAPH)
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    
    # Bold text style
    bold_style = styles.add_style('Bold Style', WD_STYLE_TYPE.CHARACTER)
    bold_font = bold_style.font
    bold_font.name = 'Arial'
    bold_font.size = Pt(11)
    bold_font.bold = True
    
    # Add title
    title = doc.add_paragraph("CORRESPONDENT BANKING AGREEMENT", style='Title Style')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("NVC FUND BANK", style='Heading Style')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    doc.add_paragraph("AGREEMENT", style='Heading Style')
    
    intro = doc.add_paragraph("THIS CORRESPONDENT BANKING AGREEMENT (the \"Agreement\") is made and entered into as of the Effective Date set forth below, by and between:", style='Normal Style')
    
    party1 = doc.add_paragraph()
    party1.add_run("NVC Fund Bank").bold = True
    doc.add_paragraph("a Supranational Sovereign Financial Institution established under the African Union Treaty framework with SWIFT code NVCFBKAU and ACH Routing Number 031176110, (hereinafter referred to as \"NVC Fund Bank\")", style='Normal Style')
    
    and_text = doc.add_paragraph("AND", style='Normal Style')
    and_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    party2 = doc.add_paragraph()
    party2.add_run("[RESPONDENT BANK NAME]").bold = True
    doc.add_paragraph("a [type of institution] organized and existing under the laws of [jurisdiction], with its principal place of business at [address], with SWIFT code [SWIFT code] (hereinafter referred to as the \"Respondent Bank\")", style='Normal Style')
    
    doc.add_paragraph("(each a \"Party\" and collectively the \"Parties\").", style='Normal Style')
    
    # Recitals
    doc.add_paragraph("RECITALS", style='Heading Style')
    
    recitals = [
        "WHEREAS, NVC Fund Bank is a Supranational Sovereign Financial Institution with over $10 trillion in balance sheet assets and maintains a global settlement infrastructure for facilitating cross-border payments, currency exchanges, and other financial services;",
        "WHEREAS, NVC Fund Bank has established a proprietary settlement system utilizing its NVC Token Stablecoin (NVCT) as the native currency for transactions, which maintains a 1:1 USD peg and is fully backed by cash and cash equivalent assets;",
        "WHEREAS, the Respondent Bank wishes to establish a correspondent banking relationship with NVC Fund Bank to access its settlement infrastructure, services, and global payment capabilities;",
        "WHEREAS, the Parties wish to set forth the terms and conditions governing their correspondent banking relationship;",
        "NOW, THEREFORE, in consideration of the mutual covenants and agreements contained herein, and for other good and valuable consideration, the receipt and sufficiency of which are hereby acknowledged, the Parties hereby agree as follows:"
    ]
    
    for recital in recitals:
        doc.add_paragraph(recital, style='Normal Style')
    
    # Definitions
    doc.add_paragraph("1. DEFINITIONS", style='Heading Style')
    
    doc.add_paragraph("In this Agreement, unless the context otherwise requires, the following terms shall have the following meanings:", style='Normal Style')
    
    definitions = [
        {"term": "\"Applicable Law\"", "definition": "means all laws, regulations, rules, orders, directives, guidelines, and standards applicable to the Parties and the services provided under this Agreement, including but not limited to banking, anti-money laundering, counter-terrorism financing, and sanctions laws."},
        {"term": "\"Business Day\"", "definition": "means a day (other than a Saturday, Sunday, or public holiday) on which banks are open for general business in the respective jurisdictions of the Parties."},
        {"term": "\"Confidential Information\"", "definition": "means any information, whether in written, oral, electronic, or other form, that is designated as confidential or that reasonably should be understood to be confidential given the nature of the information and the circumstances of disclosure."},
        {"term": "\"Correspondent Account\"", "definition": "means the account(s) established by NVC Fund Bank for the Respondent Bank pursuant to this Agreement."},
        {"term": "\"Effective Date\"", "definition": "means the date on which this Agreement is executed by both Parties."},
        {"term": "\"NVCT\"", "definition": "means the NVC Token Stablecoin, a digital currency issued by NVC Fund Bank with a 1:1 USD peg and fully backed by cash and cash equivalent assets."},
        {"term": "\"Services\"", "definition": "means the correspondent banking services to be provided by NVC Fund Bank to the Respondent Bank as set forth in this Agreement."}
    ]
    
    for definition in definitions:
        p = doc.add_paragraph(style='Normal Style')
        p.add_run(definition["term"]).bold = True
        p.add_run(" " + definition["definition"])
    
    # Scope of Services
    doc.add_paragraph("2. SCOPE OF SERVICES", style='Heading Style')
    
    doc.add_paragraph("2.1 NVC Fund Bank shall provide the following Services to the Respondent Bank:", style='Normal Style')
    
    services = [
        "Maintenance of a Correspondent Account for the Respondent Bank in one or more of the following currencies: USD, EUR, GBP, NVCT, AFD1, SFN, and other currencies as mutually agreed upon.",
        "Processing of international wire transfers through SWIFT messaging (MT103, MT202, MT760, etc.)",
        "Facilitation of foreign exchange transactions at competitive rates.",
        "Access to the NVC Fund Bank settlement system for real-time gross settlement (RTGS) capabilities.",
        "Issuance and redemption of NVCT for settlement purposes.",
        "Access to NVC Fund Bank's global payment network.",
        "Cash management services.",
        "Trade finance services, including letters of credit.",
        "Such other services as the Parties may agree upon in writing from time to time."
    ]
    
    for service in services:
        p = doc.add_paragraph(style='Normal Style')
        p.add_run("• ").bold = True
        p.add_run(service)
    
    doc.add_paragraph("2.2 The Respondent Bank shall provide NVC Fund Bank with all necessary information and documentation required for the provision of the Services, including but not limited to customer identification, transaction details, and regulatory reporting information.", style='Normal Style')
    
    doc.add_paragraph("2.3 The Parties agree to cooperate with each other in good faith to ensure the efficient provision of the Services in accordance with this Agreement and Applicable Law.", style='Normal Style')
    
    # Account Opening and Operation
    doc.add_paragraph("3. ACCOUNT OPENING AND OPERATION", style='Heading Style')
    
    account_paras = [
        "3.1 NVC Fund Bank shall open and maintain a Correspondent Account for the Respondent Bank in accordance with its standard account opening procedures and subject to satisfactory completion of all required due diligence.",
        "3.2 The Correspondent Account shall be operated in accordance with the terms and conditions set forth in this Agreement and the applicable account documentation provided by NVC Fund Bank.",
        "3.3 The Respondent Bank shall provide NVC Fund Bank with a list of authorized signatories and any specific instructions for the operation of the Correspondent Account.",
        "3.4 The Respondent Bank shall maintain sufficient funds in the Correspondent Account to cover all transactions and fees associated with the Services.",
        "3.5 NVC Fund Bank shall provide the Respondent Bank with periodic account statements and transaction confirmations in accordance with its standard practices or as otherwise agreed by the Parties."
    ]
    
    for para in account_paras:
        doc.add_paragraph(para, style='Normal Style')
    
    # Fees and Charges
    doc.add_paragraph("4. FEES AND CHARGES", style='Heading Style')
    
    fee_paras = [
        "4.1 The Respondent Bank shall pay NVC Fund Bank the fees and charges for the Services as set forth in Schedule A attached hereto, which may be amended from time to time by mutual agreement of the Parties.",
        "4.2 NVC Fund Bank shall have the right to debit the Correspondent Account for any fees, charges, or other amounts owed by the Respondent Bank under this Agreement.",
        "4.3 All fees and charges are exclusive of any applicable taxes, which shall be borne by the Respondent Bank.",
        "4.4 In addition to the fees and charges set forth in Schedule A, the Respondent Bank shall reimburse NVC Fund Bank for any out-of-pocket expenses incurred in connection with the provision of the Services, including but not limited to telecommunications, postage, and courier charges."
    ]
    
    for para in fee_paras:
        doc.add_paragraph(para, style='Normal Style')
    
    # Compliance and AML Requirements
    doc.add_paragraph("5. COMPLIANCE AND AML REQUIREMENTS", style='Heading Style')
    
    compliance_paras = [
        "5.1 Each Party shall comply with all Applicable Law in the performance of its obligations under this Agreement, including but not limited to laws and regulations relating to anti-money laundering, counter-terrorism financing, sanctions, and data protection.",
        "5.2 The Respondent Bank shall implement and maintain adequate policies, procedures, and controls to ensure compliance with Applicable Law, including but not limited to customer due diligence, transaction monitoring, and suspicious activity reporting.",
        "5.3 The Respondent Bank shall provide NVC Fund Bank with such information and documentation as NVC Fund Bank may reasonably request to fulfill its compliance obligations under Applicable Law.",
        "5.4 Each Party shall promptly notify the other Party of any material changes in its compliance policies, procedures, or controls that may affect the provision of the Services under this Agreement.",
        "5.5 NVC Fund Bank may refuse to process any transaction that it reasonably believes may violate Applicable Law or its internal policies and procedures.",
        "5.6 Each Party shall notify the other promptly if it becomes aware of any actual or suspected breach of Applicable Law in connection with the Services or this Agreement."
    ]
    
    for para in compliance_paras:
        doc.add_paragraph(para, style='Normal Style')
    
    # Term and Termination
    doc.add_paragraph("6. TERM AND TERMINATION", style='Heading Style')
    
    doc.add_paragraph("6.1 This Agreement shall commence on the Effective Date and shall continue until terminated in accordance with this Section 6.", style='Normal Style')
    
    doc.add_paragraph("6.2 Either Party may terminate this Agreement without cause by giving the other Party ninety (90) days' prior written notice.", style='Normal Style')
    
    doc.add_paragraph("6.3 Either Party may terminate this Agreement with immediate effect by giving written notice to the other Party if:", style='Normal Style')
    
    termination_reasons = [
        "The other Party commits a material breach of this Agreement and, in the case of a breach capable of remedy, fails to remedy such breach within thirty (30) days after receiving written notice to do so;",
        "The other Party becomes insolvent, is placed under administration or receivership, or undergoes any similar event;",
        "The other Party's banking license is revoked or suspended;",
        "The other Party is subject to sanctions or enforcement actions by a regulatory authority;",
        "Continuation of the Agreement would cause either Party to violate Applicable Law."
    ]
    
    for i, reason in enumerate(termination_reasons):
        p = doc.add_paragraph(style='Normal Style')
        p.add_run(chr(97 + i) + ") ").bold = True
        p.add_run(reason)
    
    doc.add_paragraph("6.4 Upon termination of this Agreement:", style='Normal Style')
    
    termination_effects = [
        "NVC Fund Bank shall prepare a final account statement and settle the Correspondent Account in accordance with the Respondent Bank's instructions, subject to any applicable legal or regulatory restrictions;",
        "Each Party shall return or destroy all Confidential Information of the other Party in its possession, except as required to be retained by Applicable Law;",
        "All outstanding obligations of the Parties shall be performed or discharged;",
        "Any accrued rights, obligations, or liabilities of the Parties shall not be affected."
    ]
    
    for i, effect in enumerate(termination_effects):
        p = doc.add_paragraph(style='Normal Style')
        p.add_run(chr(97 + i) + ") ").bold = True
        p.add_run(effect)
    
    doc.add_paragraph("6.5 The provisions of Sections 5 (Compliance and AML Requirements), 7 (Confidentiality), 8 (Data Protection), 9 (Liability and Indemnification), and 12 (Governing Law and Dispute Resolution) shall survive the termination of this Agreement.", style='Normal Style')
    
    # Additional Sections
    sections = [
        {
            "title": "7. CONFIDENTIALITY",
            "paragraphs": [
                "7.1 Each Party shall maintain the confidentiality of all Confidential Information of the other Party and shall not disclose such Confidential Information to any third party except as required by Applicable Law or with the prior written consent of the disclosing Party.",
                "7.2 Each Party shall use the Confidential Information of the other Party solely for the purpose of performing its obligations under this Agreement."
            ]
        },
        {
            "title": "8. DATA PROTECTION",
            "paragraphs": [
                "8.1 Each Party shall comply with all applicable data protection laws and regulations in the performance of its obligations under this Agreement.",
                "8.2 Each Party shall implement appropriate technical and organizational measures to protect personal data processed in connection with this Agreement."
            ]
        },
        {
            "title": "9. LIABILITY AND INDEMNIFICATION",
            "paragraphs": [
                "9.1 Each Party shall indemnify and hold harmless the other Party from and against any and all losses, damages, claims, liabilities, costs, and expenses (including reasonable attorneys' fees) arising from or in connection with any breach of this Agreement by the indemnifying Party or its negligence or willful misconduct.",
                "9.2 Neither Party shall be liable to the other Party for any indirect, special, incidental, consequential, or punitive damages arising out of or in connection with this Agreement, whether based on contract, tort, or any other legal theory.",
                "9.3 NVC Fund Bank's liability to the Respondent Bank under this Agreement shall be limited to the amount of fees paid by the Respondent Bank to NVC Fund Bank during the twelve (12) months preceding the event giving rise to such liability."
            ]
        },
        {
            "title": "10. FORCE MAJEURE",
            "paragraphs": [
                "10.1 Neither Party shall be liable for any delay or failure to perform its obligations under this Agreement to the extent that such delay or failure is caused by events beyond its reasonable control, including but not limited to acts of God, natural disasters, war, terrorism, riots, civil unrest, government actions, power failures, or telecommunications failures."
            ]
        },
        {
            "title": "11. NOTICES",
            "paragraphs": [
                "11.1 All notices and other communications under this Agreement shall be in writing and shall be delivered by hand, by courier, by registered mail, or by email to the addresses set forth below or to such other addresses as may be designated by the Parties in writing."
            ]
        },
        {
            "title": "12. GOVERNING LAW AND DISPUTE RESOLUTION",
            "paragraphs": [
                "12.1 This Agreement shall be governed by and construed in accordance with the laws of [jurisdiction], without giving effect to any choice of law or conflict of law provisions.",
                "12.2 Any dispute, controversy, or claim arising out of or in connection with this Agreement, or the breach, termination, or invalidity thereof, shall be finally settled by arbitration in accordance with the Rules of the International Chamber of Commerce by one or more arbitrators appointed in accordance with the said Rules. The place of arbitration shall be [city, country]. The language of the arbitration shall be English."
            ]
        },
        {
            "title": "13. MISCELLANEOUS",
            "paragraphs": [
                "13.1 This Agreement constitutes the entire agreement between the Parties with respect to the subject matter hereof and supersedes all prior or contemporaneous agreements, understandings, or representations, whether oral or written.",
                "13.2 This Agreement may not be amended or modified except by a written instrument signed by both Parties.",
                "13.3 Neither Party may assign this Agreement or any of its rights or obligations hereunder without the prior written consent of the other Party.",
                "13.4 If any provision of this Agreement is found to be invalid or unenforceable, the remaining provisions shall remain in full force and effect.",
                "13.5 The failure of either Party to enforce any right or provision of this Agreement shall not constitute a waiver of such right or provision.",
                "13.6 This Agreement may be executed in counterparts, each of which shall be deemed an original, but all of which together shall constitute one and the same instrument."
            ]
        }
    ]
    
    for section in sections:
        doc.add_paragraph(section["title"], style='Heading Style')
        for para in section["paragraphs"]:
            doc.add_paragraph(para, style='Normal Style')
    
    # Signature Block
    doc.add_paragraph("IN WITNESS WHEREOF, the Parties have executed this Agreement as of the Effective Date.", style='Heading Style')
    
    # Create a table for the signature block
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # NVC Fund Bank column
    table.cell(0, 0).text = "NVC FUND BANK"
    table.cell(1, 0).text = "By: ________________________"
    table.cell(2, 0).text = "Name: _____________________"
    table.cell(3, 0).text = "Title: ______________________"
    table.cell(4, 0).text = "Date: ______________________"
    
    # Respondent Bank column
    table.cell(0, 1).text = "RESPONDENT BANK"
    table.cell(1, 1).text = "By: ________________________"
    table.cell(2, 1).text = "Name: _____________________"
    table.cell(3, 1).text = "Title: ______________________"
    table.cell(4, 1).text = "Date: ______________________"
    
    # Format the signature table
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.space_after = Pt(12)
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    if j == 0 and i == 0:  # First cell in first row
                        run.font.bold = True
                    if j == 1 and i == 0:  # Second cell in first row
                        run.font.bold = True
    
    # Create the directory if it doesn't exist
    os.makedirs("static/documents", exist_ok=True)
    
    # Save the document
    docx_path = "static/documents/NVC_Fund_Bank_Correspondent_Banking_Agreement.docx"
    doc.save(docx_path)
    
    print(f"Correspondent Banking Agreement Word document generated at {docx_path}")
    return docx_path

if __name__ == "__main__":
    generate_correspondent_banking_agreement_docx()